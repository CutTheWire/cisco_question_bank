import json
import os
import random
from typing import List

import uvicorn
from docx import Document
from docx.shared import Inches
from fastapi import FastAPI, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

app = FastAPI()

# CORS 설정 (모든 출처 허용)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# 정적 파일 디렉토리 설정
app.mount("/static", StaticFiles(directory="static"), name="static")

# JSON 파일 로드
with open("questions.json", "r", encoding="utf-8") as file:
    questions_data = json.load(file)

# 요청 로깅 미들웨어
@app.middleware("http")
async def log_requests(request: Request, call_next):
    print(f"요청: {request.method} {request.url}")
    response = await call_next(request)
    print(f"응답 상태 코드: {response.status_code}")
    return response

# 요청 바디 모델 정의
class AnswerRequest(BaseModel):
    answer: List[str]

# 모든 문제를 제공하는 API 엔드포인트
@app.get("/get-questions")
async def get_questions():
    randomized_questions = random.sample(questions_data, len(questions_data))  # 문항 섞기
    questions_for_frontend = [
        {
            "index": question["index"],
            "title": question["title"],
            "options": random.sample(question["options"], len(question["options"])),  # 옵션 순서 섞기
            "image": question.get("image")
        }
        for question in randomized_questions
    ]
    return questions_for_frontend



# 사용자가 제출한 답변을 확인하는 API 엔드포인트
@app.post("/check-answer/{question_index}")
async def check_answer(question_index: int, request: AnswerRequest):
    question = next((q for q in questions_data if q["index"] == question_index), None)
    if question is None:
        raise HTTPException(status_code=404, detail="Question not found")

    correct_answer = question["answer"]

    # 정답이 문자열인 경우와 리스트인 경우를 구분하여 비교
    if isinstance(correct_answer, str):
        # 정답과 사용자의 답변을 각각 리스트로 변환하여 비교
        user_answer_list = [''.join(request.answer).replace(' ', '').replace('\n', '').strip()]
        correct_answer_list = [correct_answer.replace(' ', '').replace('\n', '').strip()]
        correct = user_answer_list == correct_answer_list
    elif isinstance(correct_answer, list):
        # 정답이 리스트인 경우, 집합으로 변환하여 순서에 상관없이 비교
        user_answer_set = set([item.replace(' ', '').replace('\n', '').strip() for item in request.answer])
        correct_answer_set = set([item.replace(' ', '').replace('\n', '').strip() for item in correct_answer])
        correct = user_answer_set == correct_answer_set
    else:
        # 예상치 못한 데이터 형식에 대한 예외 처리
        raise HTTPException(status_code=500, detail="Unexpected answer format in question data")

    # 디버깅을 위한 출력
    print(f"입력 : {user_answer_list if isinstance(correct_answer, str) else user_answer_set}")
    print(f"정답 : {correct_answer_list if isinstance(correct_answer, str) else correct_answer_set}")
    print(f"정답 여부 : {correct}")

    explanation = question.get("explanation", "")
    return {
        "correct": correct,
        "correct_answer": correct_answer,
        "explanation": explanation
    }

# 보고서 생성 및 Word 파일 반환 API
@app.post("/download-report")
async def download_report(request: Request):
    data = await request.json()
    correct_count = data.get("correctCount", 0)
    incorrect_list = data.get("incorrectList", [])

    # Word 문서 생성
    doc = Document()
    doc.add_heading("퀴즈 결과 보고서", level=1)
    doc.add_paragraph(f"맞은 개수: {correct_count}")
    doc.add_paragraph(f"틀린 개수: {len(incorrect_list)}")

    doc.add_heading("틀린 문제 목록:", level=2)
    for item in incorrect_list:
        doc.add_paragraph(f"문제: {item['title']}")
        # 이미지가 있으면 추가
        if "image" in item and item["image"]:
            image_path = os.path.join("static", item["image"])
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(2.0))
        doc.add_paragraph(f"정답: {item['correct_answer']}")
        doc.add_paragraph("")  # 빈 줄 추가

    # Word 파일로 저장
    output_path = "report.docx"
    doc.save(output_path)

    # Word 파일 응답으로 반환
    return FileResponse(output_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="report.docx")

# 루트 경로에서 index.html 반환
@app.get("/")
async def root():
    return FileResponse("static/index.html")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8111, log_level="debug")


